"""
Meeting notes generator: produces a dated .docx file for every session.

Format:
  - Header: title, date, duration
  - Summary (AI-generated)
  - Key Decisions
  - Action Items table
  - Full Transcript (timestamped)

Files saved to: ~/Documents/Meeting Notes/YYYY/MM Month/YYYY-MM-DD Title.docx
"""
from __future__ import annotations

import logging
import subprocess
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

from ..config import NotesConfig

logger = logging.getLogger(__name__)


def _add_horizontal_rule(doc: Document) -> None:
    """Add a thin horizontal line to the document."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_cell_bg(cell, hex_color: str) -> None:
    """Set table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def generate_meeting_notes(
    title: str,
    started_at: datetime,
    duration_seconds: int,
    summary: str,
    decisions: list[str],
    actions: list[dict],
    transcript: str,
    participants: list[str],
    config: NotesConfig,
    session_id: str = "",
) -> Optional[Path]:
    """
    Create a .docx meeting notes file and return its path.
    Returns None if document creation fails.
    """
    try:
        output_path = _build_output_path(title, started_at, config)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        _set_document_margins(doc)

        _add_header(doc, title, started_at, duration_seconds, participants)
        _add_horizontal_rule(doc)

        if summary:
            _add_section(doc, "SUMMARY", summary)

        if decisions:
            _add_list_section(doc, "KEY DECISIONS", decisions)

        qualifying_actions = [
            a for a in actions if a.get("confidence", 0) >= 0.5
        ]
        if qualifying_actions:
            _add_actions_table(doc, qualifying_actions)

        if transcript.strip():
            _add_transcript(doc, transcript)

        if session_id:
            _add_footer(doc, session_id)

        doc.save(str(output_path))
        logger.info(f"Meeting notes saved: {output_path}")
        return output_path

    except Exception as e:
        logger.error(f"Failed to generate meeting notes: {e}")
        return None


def open_notes(path: Path) -> None:
    """Open the Word document with the system default app."""
    subprocess.Popen(["open", str(path)])


# ── Document building helpers ─────────────────────────────────────────────────

def _build_output_path(title: str, started_at: datetime, config: NotesConfig) -> Path:
    safe_title = "".join(
        c if c.isalnum() or c in " -_" else "" for c in title
    ).strip()[:60]
    filename = f"{started_at.strftime('%Y-%m-%d')} {safe_title}.docx"

    if config.date_folders:
        folder = config.output_path / started_at.strftime("%Y") / started_at.strftime("%m %B")
    else:
        folder = config.output_path

    return folder / filename


def _set_document_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)


def _add_header(
    doc: Document,
    title: str,
    started_at: datetime,
    duration_seconds: int,
    participants: list[str],
) -> None:
    # Title
    heading = doc.add_heading(title.upper(), level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # Metadata line
    h = duration_seconds // 3600
    m = (duration_seconds % 3600) // 60
    s = duration_seconds % 60
    duration_str = f"{h:02d}:{m:02d}:{s:02d}"
    date_str = started_at.strftime("%B %d, %Y · %I:%M %p")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = meta.add_run(f"{date_str}  ·  Duration: {duration_str}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    if participants:
        pline = doc.add_paragraph()
        r = pline.add_run(f"Participants: {', '.join(participants)}")
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        r.font.italic = True

    watermark = doc.add_paragraph()
    wr = watermark.add_run("Recorded by nudge")
    wr.font.size = Pt(8)
    wr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.add_paragraph()  # spacer


def _add_section(doc: Document, heading_text: str, body: str) -> None:
    h = doc.add_heading(heading_text, level=2)
    for run in h.runs:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    p = doc.add_paragraph(body)
    p.paragraph_format.space_after = Pt(12)
    doc.add_paragraph()


def _add_list_section(doc: Document, heading_text: str, items: list[str]) -> None:
    h = doc.add_heading(heading_text, level=2)
    for run in h.runs:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    doc.add_paragraph()


def _add_actions_table(doc: Document, actions: list[dict]) -> None:
    h = doc.add_heading("ACTION ITEMS", level=2)
    for run in h.runs:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["#", "Task", "Owner", "Due"]):
        cell.text = text
        _set_cell_bg(cell, "1A1A2E")
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)

    # Data rows
    for i, action in enumerate(actions, start=1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = action.get("task", "")
        row[2].text = action.get("assignee") or "—"
        row[3].text = action.get("deadline") or "—"

        bg = "F0F4FF" if i % 2 == 0 else "FFFFFF"
        for cell in row:
            _set_cell_bg(cell, bg)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    # Column widths
    for row in table.rows:
        row.cells[0].width = Inches(0.4)
        row.cells[1].width = Inches(3.5)
        row.cells[2].width = Inches(1.2)
        row.cells[3].width = Inches(1.2)

    doc.add_paragraph()


def _add_transcript(doc: Document, transcript: str) -> None:
    _add_horizontal_rule(doc)
    h = doc.add_heading("FULL TRANSCRIPT", level=2)
    for run in h.runs:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # Split into paragraphs for readability
    for line in transcript.split("\n"):
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def _add_footer(doc: Document, session_id: str) -> None:
    """Add session ID to document footer for traceability."""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"nudge session: {session_id}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
